import asyncio
import os
import unittest
from io import BytesIO
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.main import app, auth_service
from app.services.document_service import DocumentService, MAX_UPLOAD_BYTES
from app.services.postgres_vector_store import PostgresVectorStore


def make_pdf(text: str | None = None) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_ref}
                )
            }
        )
        stream = DecodedStreamObject()
        safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream.set_data(
            f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode("ascii")
        )
        page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_docx() -> bytes:
    document = DocxDocument()
    document.add_heading("Orchid Protocol", level=1)
    document.add_paragraph("Orchid access requires manager approval and a security key.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class M7DocumentApiTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.client = TestClient(app)
        username = f"m7-admin-{uuid4().hex}"
        auth_service.create_user(username, "M7-document-pass-123!", "admin")
        response = cls.client.post(
            "/auth/token",
            data={"username": username, "password": "M7-document-pass-123!"},
        )
        cls.headers = {
            "Authorization": f"Bearer {response.json()['access_token']}"
        }

    def setUp(self) -> None:
        self._delete_all_documents()

    def tearDown(self) -> None:
        self._delete_all_documents()

    def _delete_all_documents(self) -> None:
        for document in self.client.get(
            "/documents", headers=self.headers
        ).json():
            self.client.delete(
                f"/documents/{document['document_id']}", headers=self.headers
            )

    def _upload(self, filename: str, data: bytes, content_type: str):
        return self.client.post(
            "/documents/upload-file",
            files={"file": (filename, data, content_type)},
            headers=self.headers,
        )

    def test_txt_pdf_docx_sources_and_delete(self) -> None:
        text = self.client.post(
            "/documents/upload",
            json={"filename": "retention.txt", "content": "Records are retained for seven years."},
            headers=self.headers,
        )
        pdf = self._upload(
            "travel-policy.pdf",
            make_pdf("Quartz travel reimbursement requires invoice and itinerary."),
            "application/pdf",
        )
        docx = self._upload(
            "access-policy.docx",
            make_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        for response, file_type in ((text, "txt"), (pdf, "pdf"), (docx, "docx")):
            self.assertEqual(response.status_code, 200, response.text)
            self.assertEqual(response.json()["file_type"], file_type)

        pdf_chunks = self.client.get(
            f"/documents/{pdf.json()['document_id']}/chunks", headers=self.headers
        ).json()
        self.assertEqual(pdf_chunks[0]["file_type"], "pdf")
        self.assertEqual(pdf_chunks[0]["page_number"], 1)

        pdf_answer = self.client.post(
            "/qa/ask",
            json={"question": "Quartz travel reimbursement invoice", "top_k": 3},
            headers=self.headers,
        ).json()
        pdf_source = next(
            source for source in pdf_answer["sources"]
            if source["filename"] == "travel-policy.pdf"
        )
        self.assertEqual(pdf_source["page_number"], 1)
        self.assertIn("页码：1", pdf_answer["prompt"])

        docx_answer = self.client.post(
            "/qa/ask",
            json={"question": "Orchid manager approval security key", "top_k": 3},
            headers=self.headers,
        ).json()
        docx_source = next(
            source for source in docx_answer["sources"]
            if source["filename"] == "access-policy.docx"
        )
        self.assertIn("Orchid Protocol", docx_source["section"])
        self.assertIn("段落", docx_source["section"])

        for response in (text, pdf, docx):
            document_id = response.json()["document_id"]
            deleted = self.client.delete(
                f"/documents/{document_id}", headers=self.headers
            )
            self.assertEqual(deleted.status_code, 200)
            self.assertEqual(
                self.client.get(
                    f"/documents/{document_id}/chunks", headers=self.headers
                ).status_code,
                404,
            )

    def test_invalid_empty_large_scanned_and_duplicate_files(self) -> None:
        cases = [
            ("data.csv", b"value", "text/csv"),
            ("empty.txt", b"", "text/plain"),
            ("broken.pdf", b"not a PDF", "application/pdf"),
            ("broken.docx", b"not a DOCX", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            ("scanned.pdf", make_pdf(), "application/pdf"),
            ("large.txt", b"x" * (MAX_UPLOAD_BYTES + 1), "text/plain"),
        ]
        for filename, data, content_type in cases:
            with self.subTest(filename=filename):
                response = self._upload(filename, data, content_type)
                self.assertEqual(response.status_code, 400, response.text)
        self.assertIn("OCR", self._upload("scan.pdf", make_pdf(), "application/pdf").text)

        first = self._upload("first.txt", b"duplicate policy text", "text/plain")
        duplicate = self._upload("second.txt", b"duplicate policy text", "text/plain")
        self.assertEqual(first.status_code, 200)
        self.assertEqual(duplicate.status_code, 409)


class LocalEmbeddings:
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self.embed_query(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        return [1.0] + [0.0] * 1023

    async def aembed_documents(self, texts: list[str]) -> list[list[float]]:
        return self.embed_documents(texts)

    async def aembed_query(self, text: str) -> list[float]:
        return self.embed_query(text)


@unittest.skipUnless(
    os.getenv("RUN_POSTGRES_INTEGRATION") == "1",
    "set RUN_POSTGRES_INTEGRATION=1 to test M7 PostgreSQL metadata",
)
class M7PostgresTest(unittest.TestCase):
    def test_source_metadata_persists_and_deletes_without_paid_apis(self) -> None:
        embedder = LocalEmbeddings()
        store = PostgresVectorStore(os.environ["DATABASE_URL"], embedder, 1024)
        document_id = None
        filename = f"m7-{uuid4().hex}.pdf"
        try:
            document = DocumentService(store, embedder).add_file_document(
                filename,
                make_pdf(f"M7 persistence marker {uuid4().hex}."),
            )
            document_id = document["document_id"]
            chunks = store.get_chunks(document_id)
            self.assertEqual(chunks[0].file_type, "pdf")
            self.assertEqual(chunks[0].page_number, 1)
            self.assertTrue(store.delete_document(document_id))
            document_id = None
            self.assertEqual(store.get_chunks(document["document_id"]), [])
        finally:
            if document_id:
                store.delete_document(document_id)
            for item in store.list_documents():
                if item["filename"] == filename:
                    store.delete_document(item["document_id"])
            asyncio.run(store.close())


if __name__ == "__main__":
    unittest.main()
